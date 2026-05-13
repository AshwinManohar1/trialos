import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ──────────────────────────────────────────────
# Placeholder map builder
# ──────────────────────────────────────────────

def _build_variable_map(study, drug_profile, pk, ai_sections: dict = None) -> dict:
    """Build the placeholder → value mapping, including AI-generated sections."""
    regulatory_targets_str = (
        ", ".join(drug_profile.regulatory_targets)
        if drug_profile.regulatory_targets
        else ""
    )

    timepoints = pk.pk_sampling_timepoints or []
    timepoints_str = ", ".join(
        (f"{t:g}h" if t != int(t) else f"{int(t)}h") for t in timepoints
    )

    ambulatory = ", ".join(pk.ambulatory_visits) if pk.ambulatory_visits else "None"
    ai = ai_sections or {}

    # Format inclusion/exclusion as numbered text
    def _list_to_text(lst):
        if not lst:
            return ""
        lines = []
        for i, item in enumerate(lst, 1):
            # Strip leading number if AI already added it
            text = str(item).strip()
            if text and text[0].isdigit() and text[1] in ".):":
                lines.append(text)
            else:
                lines.append(f"{i}. {text}")
        return "\n".join(lines)

    return {
        # ── Basic identifiers ──
        "{{STUDY_ID}}": study.id or "",
        "{{VERSION}}": "1.0",
        "{{DATE}}": study.created_at.strftime("%d %B %Y") if study.created_at else "",
        "{{STUDY_TITLE}}": ai.get("study_title", f"Bioequivalence Study of {drug_profile.drug_name} {drug_profile.dose}"),

        # ── Drug & sponsor ──
        "{{DRUG_NAME}}": drug_profile.drug_name or "",
        "{{DOSE}}": drug_profile.dose or "",
        "{{FORMULATION}}": drug_profile.formulation or "",
        "{{ROUTE}}": drug_profile.route or "Oral",
        "{{REFERENCE_PRODUCT}}": drug_profile.reference_product or "",
        "{{REFERENCE_COUNTRY}}": drug_profile.reference_country or "",
        "{{SPONSOR_NAME}}": drug_profile.sponsor_name or "",
        "{{SPONSOR_COUNTRY}}": drug_profile.sponsor_country or "",
        "{{REGULATORY_TARGETS}}": regulatory_targets_str,
        "{{TARGET_SUBJECTS}}": str(drug_profile.target_subjects or ""),
        "{{SPECIAL_INSTRUCTIONS}}": drug_profile.special_instructions or "",

        # ── PK parameters ──
        "{{HALF_LIFE}}": str(pk.half_life_hours or ""),
        "{{TMAX}}": str(pk.tmax_hours or ""),
        "{{ABSORPTION_CLASS}}": pk.absorption_class or "",
        "{{WASHOUT_DAYS}}": str(pk.washout_days or ""),
        "{{CONFINEMENT_HOURS}}": str(pk.confinement_hours or ""),
        "{{PK_SAMPLING_SCHEDULE}}": timepoints_str,
        "{{AMBULATORY_VISITS}}": ambulatory,
        "{{POSTURE_RESTRICTION}}": pk.posture_restriction or "No specific posture restriction",
        "{{SAMPLE_SIZE_RECOMMENDED}}": str(pk.sample_size_recommended or ""),
        "{{SAMPLE_SIZE_BASIS}}": pk.sample_size_basis or "",
        "{{INTRASUBJECT_CV}}": f"{pk.intrasubject_cv}%" if pk.intrasubject_cv else "",

        # ── AI-generated narrative sections ──
        "{{STUDY_DESIGN_NARRATIVE}}": ai.get("study_design_narrative", ""),
        "{{PRIMARY_OBJECTIVE}}": ai.get("primary_objective", ""),
        "{{SECONDARY_OBJECTIVES}}": "\n".join(f"• {o}" for o in ai.get("secondary_objectives", [])),
        "{{INCLUSION_CRITERIA}}": _list_to_text(ai.get("inclusion_criteria", [])),
        "{{EXCLUSION_CRITERIA}}": _list_to_text(ai.get("exclusion_criteria", [])),
        "{{WITHDRAWAL_CRITERIA}}": _list_to_text(ai.get("withdrawal_criteria", [])),
        "{{CONFINEMENT_SCHEDULE}}": ai.get("confinement_schedule", ""),
        "{{DOSE_ADMINISTRATION}}": ai.get("dose_administration", ""),
        "{{PK_SAMPLING_PROCEDURE}}": ai.get("pk_sampling_procedure", ""),
        "{{SAFETY_ASSESSMENTS}}": ai.get("safety_assessments", ""),
        "{{DIET_RESTRICTIONS}}": ai.get("diet_restrictions", ""),
        "{{STATISTICAL_ANALYSIS}}": ai.get("statistical_analysis", ""),
        "{{STOPPING_CRITERIA}}": _list_to_text(ai.get("stopping_criteria", [])),
        "{{ADVERSE_EVENT_MANAGEMENT}}": ai.get("adverse_event_management", ""),
        "{{REFERENCES}}": "\n".join(
            f"{i+1}. {r}" for i, r in enumerate(ai.get("references", []))
        ),
    }


# ──────────────────────────────────────────────
# Template placeholder replacement
# ──────────────────────────────────────────────

def _replace_in_paragraph(paragraph, var_map: dict):
    """Replace placeholders in a paragraph, preserving formatting."""
    for run in paragraph.runs:
        for placeholder, value in var_map.items():
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)

    full_text = "".join(run.text for run in paragraph.runs)
    for placeholder, value in var_map.items():
        if placeholder in full_text:
            if paragraph.runs:
                paragraph.runs[0].text = full_text.replace(placeholder, value)
                for run in paragraph.runs[1:]:
                    run.text = ""
                full_text = paragraph.runs[0].text


def _has_placeholders(doc: Document) -> bool:
    for para in doc.paragraphs:
        if "{{" in para.text:
            return True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "{{" in para.text:
                        return True
    return False


# ──────────────────────────────────────────────
# Full protocol DOCX generator (no template)
# ──────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_section_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _add_body_text(doc: Document, text: str):
    if not text:
        return
    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())


def _add_numbered_list(doc: Document, items):
    """Add items as a numbered list, handling pre-numbered strings too."""
    if not items:
        return
    if isinstance(items, str):
        items = [l for l in items.split("\n") if l.strip()]
    for item in items:
        text = str(item).strip()
        if not text:
            continue
        p = doc.add_paragraph(style="List Number")
        # Strip leading "1. " etc if already numbered
        if len(text) > 2 and text[0].isdigit() and text[1] in ".)":
            text = text[2:].strip()
        elif len(text) > 3 and text[0].isdigit() and text[1].isdigit() and text[2] in ".)":
            text = text[3:].strip()
        p.add_run(text)


def _add_bullet_list(doc: Document, items):
    if not items:
        return
    if isinstance(items, str):
        items = [l.lstrip("•- ") for l in items.split("\n") if l.strip()]
    for item in items:
        text = str(item).strip().lstrip("•- ")
        if text:
            doc.add_paragraph(text, style="List Bullet")


def _add_two_col_table(doc: Document, rows: list, header: tuple = None):
    """Add a simple 2-column table."""
    n_rows = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=n_rows, cols=2)
    table.style = "Table Grid"

    row_idx = 0
    if header:
        cells = table.rows[row_idx].cells
        cells[0].text = header[0]
        cells[1].text = header[1]
        for cell in cells:
            _set_cell_bg(cell, "1B2A3B")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)
        row_idx += 1

    for label, value in rows:
        cells = table.rows[row_idx].cells
        cells[0].text = str(label)
        cells[1].text = str(value) if value is not None else ""
        for para in cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
        for para in cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
        row_idx += 1

    return table


def _generate_full_protocol_docx(doc: Document, study, drug_profile, pk, ai: dict):
    """Generate a complete, professional BE protocol document."""
    reg_str = ", ".join(drug_profile.regulatory_targets or []) or "US FDA"
    timepoints = pk.pk_sampling_timepoints or []
    tp_str = ", ".join((f"{t:g}h" if t != int(t) else f"{int(t)}h") for t in timepoints)
    ambulatory = ", ".join(pk.ambulatory_visits or []) or "None"
    study_title = ai.get("study_title", f"Bioequivalence Study of {drug_profile.drug_name} {drug_profile.dose}")

    # ── Cover Page ──────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("STUDY PROTOCOL")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x3B)

    doc.add_paragraph()

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = title2.add_run(study_title)
    run2.bold = True
    run2.font.size = Pt(13)

    doc.add_paragraph()

    meta_rows = [
        ("Protocol Number", study.id),
        ("Version", "1.0"),
        ("Date", study.created_at.strftime("%d %B %Y") if study.created_at else ""),
        ("Sponsor", f"{drug_profile.sponsor_name}, {drug_profile.sponsor_country}"),
        ("Regulatory Targets", reg_str),
        ("Phase", "Phase I — Bioequivalence"),
        ("Study Design", "Randomized, Open-Label, 2-Period, 2-Treatment, 2-Sequence Crossover"),
    ]
    _add_two_col_table(doc, meta_rows)

    doc.add_page_break()

    # ── Synopsis ────────────────────────────────
    _add_section_heading(doc, "SYNOPSIS", 1)

    synopsis_rows = [
        ("Study Title", study_title),
        ("Protocol Number", study.id),
        ("Phase", "Phase I — Bioequivalence"),
        ("Study Design", "Randomized, Open-Label, 2-Period, 2-Treatment, 2-Sequence, Single-Dose Crossover"),
        ("Test Product", f"{drug_profile.drug_name} {drug_profile.dose} {drug_profile.formulation}"),
        ("Reference Product", f"{drug_profile.reference_product} ({drug_profile.reference_country})"),
        ("Regulatory Submission", reg_str),
        ("Sponsor", f"{drug_profile.sponsor_name}, {drug_profile.sponsor_country}"),
        ("Target Subjects", str(drug_profile.target_subjects or pk.sample_size_recommended)),
        ("Primary Objective", ai.get("primary_objective", "To evaluate the bioequivalence of the test and reference formulations.")),
        ("Administration", f"Single oral dose, {drug_profile.route}"),
        ("Fasting Condition", "Fasting (≥10 hours pre-dose)" if "fasting" in (drug_profile.special_instructions or "").lower() else "As specified in protocol"),
        ("Washout Period", f"{pk.washout_days} days"),
        ("Confinement", f"{pk.confinement_hours} hours"),
        ("PK Sampling", tp_str),
        ("Ambulatory Visits", ambulatory),
        ("Primary PK Parameters", "AUC0-t, AUC0-inf, Cmax"),
        ("BE Acceptance Criteria", "90% CI of GMR within 80.00–125.00%"),
        ("Sample Size", f"n = {pk.sample_size_recommended} ({pk.sample_size_basis})"),
    ]
    _add_two_col_table(doc, synopsis_rows, header=("Parameter", "Details"))

    doc.add_page_break()

    # ── 1. Introduction ──────────────────────────
    _add_section_heading(doc, "1. INTRODUCTION AND BACKGROUND", 1)
    _add_body_text(doc, ai.get("study_design_narrative", ""))
    if drug_profile.special_instructions:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Special Considerations: ")
        run.bold = True
        p.add_run(drug_profile.special_instructions)

    # ── 2. Objectives ────────────────────────────
    doc.add_page_break()
    _add_section_heading(doc, "2. STUDY OBJECTIVES", 1)

    _add_section_heading(doc, "2.1 Primary Objective", 2)
    doc.add_paragraph(ai.get("primary_objective", ""))

    _add_section_heading(doc, "2.2 Secondary Objectives", 2)
    _add_bullet_list(doc, ai.get("secondary_objectives", []))

    # ── 3. Study Design ──────────────────────────
    doc.add_page_break()
    _add_section_heading(doc, "3. STUDY DESIGN", 1)
    _add_body_text(doc, ai.get("study_design_narrative", ""))

    _add_section_heading(doc, "3.1 Study Schematic", 2)
    design_rows = [
        ("Design", "Randomized, Open-Label, 2-Period, 2-Sequence, Single-Dose Crossover"),
        ("Periods", "2 treatment periods"),
        ("Sequences", f"Sequence 1: Test → Reference | Sequence 2: Reference → Test"),
        ("Washout Period", f"{pk.washout_days} days (≥ 5 × t½ = {round(pk.half_life_hours * 5, 1)}h)"),
        ("Confinement", f"Approximately {pk.confinement_hours} hours post-dose"),
        ("Total Duration", f"Approximately {pk.washout_days + 3} days per subject (including both periods)"),
        ("Fasting", "≥10 hours pre-dose; 4 hours post-dose (water allowed)"),
        ("Ambulatory Follow-up", ambulatory),
    ]
    _add_two_col_table(doc, design_rows, header=("Parameter", "Details"))

    # ── 4. Subject Selection ─────────────────────
    doc.add_page_break()
    _add_section_heading(doc, "4. SUBJECT SELECTION", 1)

    _add_section_heading(doc, "4.1 Inclusion Criteria", 2)
    doc.add_paragraph("Subjects must meet ALL of the following criteria to be eligible for participation:")
    _add_numbered_list(doc, ai.get("inclusion_criteria", []))

    _add_section_heading(doc, "4.2 Exclusion Criteria", 2)
    doc.add_paragraph("Subjects meeting ANY of the following criteria will be excluded:")
    _add_numbered_list(doc, ai.get("exclusion_criteria", []))

    _add_section_heading(doc, "4.3 Withdrawal Criteria", 2)
    doc.add_paragraph("Subjects may be withdrawn from the study if any of the following occur:")
    _add_numbered_list(doc, ai.get("withdrawal_criteria", []))

    # ── 5. Study Procedures ──────────────────────
    doc.add_page_break()
    _add_section_heading(doc, "5. STUDY PROCEDURES", 1)

    _add_section_heading(doc, "5.1 Screening", 2)
    doc.add_paragraph(
        "Screening will be conducted within 21 days prior to the first dose. "
        "Subjects will undergo a full medical history, physical examination, vital signs, "
        "12-lead ECG, haematology, clinical chemistry, urinalysis, urine drug screen, "
        "alcohol breath test, and pregnancy test (for females)."
    )

    _add_section_heading(doc, "5.2 Confinement Schedule", 2)
    _add_body_text(doc, ai.get("confinement_schedule", ""))

    _add_section_heading(doc, "5.3 Dose Administration", 2)
    _add_body_text(doc, ai.get("dose_administration", ""))

    _add_section_heading(doc, "5.4 Dietary Restrictions", 2)
    _add_body_text(doc, ai.get("diet_restrictions", ""))

    # ── 6. PK Blood Sampling ─────────────────────
    doc.add_page_break()
    _add_section_heading(doc, "6. PHARMACOKINETIC BLOOD SAMPLING", 1)

    _add_section_heading(doc, "6.1 Sampling Schedule", 2)
    doc.add_paragraph(
        f"Blood samples will be collected at the following timepoints relative to dose: {tp_str}"
    )

    # Sampling timepoints table
    if timepoints:
        in_house = [t for t in timepoints if t <= pk.confinement_hours]
        ambul = [t for t in timepoints if t > pk.confinement_hours]

        _add_section_heading(doc, "In-Clinic Samples", 3)
        tp_table_rows = [(f"{t:g}h", "4–6 mL EDTA whole blood") for t in in_house]
        _add_two_col_table(doc, tp_table_rows, header=("Timepoint (post-dose)", "Volume & Tube Type"))

        if ambul:
            _add_section_heading(doc, "Ambulatory Samples", 3)
            amb_rows = [(f"{t:g}h", "4–6 mL EDTA — subject returns to clinic") for t in ambul]
            _add_two_col_table(doc, amb_rows, header=("Timepoint (post-dose)", "Collection"))

    _add_section_heading(doc, "6.2 Sample Collection and Processing", 2)
    _add_body_text(doc, ai.get("pk_sampling_procedure", ""))

    # ── 7. Safety Assessments ────────────────────
    doc.add_page_break()
    _add_section_heading(doc, "7. SAFETY ASSESSMENTS", 1)
    _add_body_text(doc, ai.get("safety_assessments", ""))

    if pk.safety_flags:
        _add_section_heading(doc, "7.1 Drug-Specific Safety Monitoring", 2)
        for flag in pk.safety_flags:
            _add_section_heading(doc, flag.get("type", "Safety Flag"), 3)
            doc.add_paragraph(flag.get("description", ""))
            reqs = flag.get("requirements", [])
            if reqs:
                _add_bullet_list(doc, reqs)

    _add_section_heading(doc, "7.2 Adverse Event Management", 2)
    _add_body_text(doc, ai.get("adverse_event_management", ""))

    _add_section_heading(doc, "7.3 Stopping Criteria", 2)
    _add_numbered_list(doc, ai.get("stopping_criteria", []))

    # ── 8. Statistical Analysis ──────────────────
    doc.add_page_break()
    _add_section_heading(doc, "8. STATISTICAL ANALYSIS PLAN", 1)
    _add_body_text(doc, ai.get("statistical_analysis", ""))

    stat_rows = [
        ("Primary PK Parameters", "AUC0-t, AUC0-inf, Cmax"),
        ("Secondary PK Parameters", "t½, Tmax, Kel"),
        ("Statistical Model", "ANOVA with factors: sequence, period, treatment, subjects within sequence"),
        ("Point Estimate", "Geometric Mean Ratio (Test/Reference)"),
        ("Confidence Interval", "90% CI using ln-transformed PK parameters"),
        ("BE Acceptance Criteria", "90% CI within 80.00–125.00%"),
        ("Intrasubject CV", f"{pk.intrasubject_cv}%" if pk.intrasubject_cv else "As determined"),
        ("Sample Size", f"n = {pk.sample_size_recommended}"),
        ("Sample Size Basis", pk.sample_size_basis or ""),
        ("Statistical Software", "Phoenix WinNonlin v8.x or higher"),
        ("Analysis Population", "Per-protocol (PP) population; sensitivity analysis on ITT"),
    ]
    _add_two_col_table(doc, stat_rows, header=("Parameter", "Details"))

    # ── 9. References ─────────────────────────────
    doc.add_page_break()
    _add_section_heading(doc, "9. REFERENCES", 1)
    refs = ai.get("references", [])
    if pk.source_references:
        # Merge AI refs with PK lookup refs
        all_refs = list(dict.fromkeys(refs + pk.source_references))
    else:
        all_refs = refs
    for i, ref in enumerate(all_refs, 1):
        doc.add_paragraph(f"{i}. {ref}")

    # ── Footer note ───────────────────────────────
    doc.add_page_break()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        f"Generated by TrialOS Clinical Intelligence Platform\n"
        f"Protocol: {study.id} | Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}\n"
        "This document is a draft and must be reviewed by a qualified clinical scientist before submission."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ──────────────────────────────────────────────
# Main entry point
# ──────────────────────────────────────────────

def fill_protocol_template_bytes(
    template_bytes: bytes,
    study,
    drug_profile,
    pk,
    ai_sections: dict = None,
) -> bytes:
    """
    Fill template from bytes. If template has {{placeholders}}, replaces them.
    If no placeholders, generates a full structured protocol document.
    Returns bytes of the filled/generated DOCX.
    """
    doc = Document(io.BytesIO(template_bytes))
    var_map = _build_variable_map(study, drug_profile, pk, ai_sections)
    has_placeholders = _has_placeholders(doc)

    if has_placeholders:
        for para in doc.paragraphs:
            _replace_in_paragraph(para, var_map)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_in_paragraph(para, var_map)
        for section in doc.sections:
            for para in section.header.paragraphs:
                _replace_in_paragraph(para, var_map)
            for para in section.footer.paragraphs:
                _replace_in_paragraph(para, var_map)
    else:
        # No placeholders in template — generate full protocol
        _generate_full_protocol_docx(doc, study, drug_profile, pk, ai_sections or {})

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


def fill_protocol_template(template_path: str, study, drug_profile, pk, ai_sections: dict = None) -> str:
    """Legacy file-based wrapper."""
    import os

    with open(template_path, "rb") as f:
        template_bytes = f.read()

    filled_bytes = fill_protocol_template_bytes(template_bytes, study, drug_profile, pk, ai_sections)

    generated_dir = os.getenv("GENERATED_DIR", "./generated")
    os.makedirs(generated_dir, exist_ok=True)

    base_name = os.path.basename(template_path)
    name_no_ext = os.path.splitext(base_name)[0]
    out_filename = f"{study.id}_{name_no_ext}_filled.docx"
    out_path = os.path.join(generated_dir, out_filename)

    with open(out_path, "wb") as f:
        f.write(filled_bytes)

    return out_path


def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        texts.append(para.text)
    return "\n".join(texts)


def extract_text_from_docx_bytes(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        texts.append(para.text)
    return "\n".join(texts)
