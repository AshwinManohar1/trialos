"""
Generate a realistic BE protocol Word template for Meridian Clinical Research Centre.
Run from the backend/ directory:
    python3 scripts/generate_template.py
Output: templates/meridian_be_protocol_template_v1.docx
"""

import os
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")
os.makedirs(OUTPUT_DIR, exist_ok=True)
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "meridian_be_protocol_template_v1.docx")

COMPANY = "Meridian Clinical Research Centre"
COMPANY_SHORT = "MCRC"
COMPANY_ADDR = "12 Pharma Park, Sector 18, Navi Mumbai – 400 705, India"
COMPANY_PHONE = "+91 22 6100 4400"
COMPANY_EMAIL = "protocols@meridiancrc.com"
NAVY = RGBColor(0x1E, 0x2A, 0x3A)
TEAL = RGBColor(0x4A, 0x8B, 0x7B)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, navy=True):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if navy else TEAL
        run.font.bold = True
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    return h


def add_para(doc, text="", bold=False, italic=False, size=10, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p


def add_placeholder_row(doc, label, placeholder, note=""):
    """Add a labeled field row with placeholder."""
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.columns[0].width = Inches(2.0)
    t.columns[1].width = Inches(2.5)
    t.columns[2].width = Inches(2.0)
    row = t.rows[0]
    row.cells[0].text = label
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[1].text = placeholder
    row.cells[1].paragraphs[0].runs[0].font.color.rgb = TEAL
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[1].paragraphs[0].runs[0].italic = True
    row.cells[2].text = note
    row.cells[2].paragraphs[0].runs[0].font.size = Pt(8)
    row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    doc.add_paragraph()
    return t


def build_template():
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # ── TITLE PAGE ────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(COMPANY)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("STUDY PROTOCOL")
    run2.font.size = Pt(22)
    run2.font.bold = True
    run2.font.color.rgb = TEAL

    doc.add_paragraph()

    # Protocol info box
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = "Table Grid"
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_rows = [
        ("Protocol Number:", "{{STUDY_ID}}"),
        ("Title:", "A Randomized, Open-Label, Single-Dose, Two-Period Crossover Bioequivalence Study of {{DRUG_NAME}} {{DOSE}} {{FORMULATION}} ({{SPONSOR_NAME}}) versus {{REFERENCE_PRODUCT}} ({{REFERENCE_COUNTRY}}) in Healthy Adult Subjects Under Fasting Conditions"),
        ("Version:", "{{VERSION}}"),
        ("Date:", "{{DATE}}"),
        ("Sponsor:", "{{SPONSOR_NAME}}, {{SPONSOR_COUNTRY}}"),
        ("Regulatory Submission:", "{{REGULATORY_TARGETS}}"),
    ]
    for i, (label, value) in enumerate(info_rows):
        row = info_table.rows[i]
        set_cell_bg(row.cells[0], "1E2A3A")
        lrun = row.cells[0].paragraphs[0].add_run(label)
        lrun.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        lrun.font.bold = True
        lrun.font.size = Pt(9)
        vrun = row.cells[1].paragraphs[0].add_run(value)
        vrun.font.color.rgb = TEAL
        vrun.font.italic = True
        vrun.font.size = Pt(9)

    doc.add_page_break()

    # ── CONFIDENTIALITY STATEMENT ─────────────────────────────────────────────
    add_heading(doc, "CONFIDENTIALITY STATEMENT", level=1)
    add_para(doc, (
        f"This protocol is the confidential property of {COMPANY} ({COMPANY_SHORT}) and the Sponsor. "
        "It must not be reproduced, disclosed, or used without prior written consent. "
        "This document is intended solely for the individuals involved in the conduct of this study."
    ))
    doc.add_paragraph()

    # ── SYNOPSIS ─────────────────────────────────────────────────────────────
    add_heading(doc, "SYNOPSIS", level=1)
    syn_table = doc.add_table(rows=14, cols=2)
    syn_table.style = "Table Grid"
    syn_rows = [
        ("Protocol Number", "{{STUDY_ID}}"),
        ("Title", "Bioequivalence Study of {{DRUG_NAME}} {{DOSE}} {{FORMULATION}} versus {{REFERENCE_PRODUCT}}"),
        ("Phase", "Bioequivalence (Phase 1)"),
        ("Study Design", "Randomized, open-label, single-dose, two-period, two-sequence crossover"),
        ("Test Product", "{{DRUG_NAME}} {{DOSE}} {{FORMULATION}} — Manufacturer: {{SPONSOR_NAME}}"),
        ("Reference Product", "{{REFERENCE_PRODUCT}} — Country: {{REFERENCE_COUNTRY}}"),
        ("Route of Administration", "Oral"),
        ("Study Population", "Healthy adult male and female subjects, aged 18–55 years"),
        ("Number of Subjects", "{{TARGET_SUBJECTS}} (planned); minimum {{SAMPLE_SIZE_RECOMMENDED}} evaluable"),
        ("Washout Period", "{{WASHOUT_DAYS}} days (≥5 × t½ of {{DRUG_NAME}})"),
        ("Confinement", "{{CONFINEMENT_HOURS}} hours in-house per period"),
        ("Primary PK Parameters", "AUC₀‑t, AUC₀‑∞, Cmax"),
        ("Bioequivalence Criterion", "90% CI of Test/Reference GMR within 80.00–125.00%"),
        ("Regulatory Target(s)", "{{REGULATORY_TARGETS}}"),
    ]
    for i, (label, value) in enumerate(syn_rows):
        row = syn_table.rows[i]
        lrun = row.cells[0].paragraphs[0].add_run(label)
        lrun.font.bold = True
        lrun.font.size = Pt(9)
        set_cell_bg(row.cells[0], "EBF5F2")
        vrun = row.cells[1].paragraphs[0].add_run(value)
        vrun.font.size = Pt(9)
        if "{{" in value:
            vrun.font.color.rgb = TEAL
            vrun.font.italic = True

    doc.add_page_break()

    # ── TABLE OF CONTENTS ─────────────────────────────────────────────────────
    add_heading(doc, "TABLE OF CONTENTS", level=1)
    toc_items = [
        "1. Introduction and Background",
        "2. Study Objectives",
        "3. Investigational Medicinal Products",
        "4. Study Design",
        "5. Subject Selection",
        "6. Study Procedures",
        "7. Pharmacokinetic Sampling",
        "8. Pharmacokinetic Analysis",
        "9. Statistical Analysis",
        "10. Safety Monitoring",
        "11. Data Management",
        "12. Ethics and Regulatory Compliance",
        "13. Quality Assurance",
        "14. References",
        "Appendix A: Schedule of Events",
        "Appendix B: Subject Informed Consent Form",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ── SECTION 1: INTRODUCTION ───────────────────────────────────────────────
    add_heading(doc, "1. INTRODUCTION AND BACKGROUND", level=1)
    add_para(doc, (
        "{{DRUG_NAME}} ({{DOSE}}, {{FORMULATION}}) is a pharmaceutical agent proposed for marketing authorization "
        "by {{SPONSOR_NAME}}. This study is designed to demonstrate bioequivalence between the test formulation "
        "and the reference product {{REFERENCE_PRODUCT}} as required by the regulatory agencies of {{REGULATORY_TARGETS}}."
    ))
    add_para(doc, (
        "The pharmacokinetic profile of {{DRUG_NAME}} is characterized by a half-life (t½) of approximately "
        "{{HALF_LIFE}} hours and a time to peak concentration (Tmax) of approximately {{TMAX}} hours. "
        "The intrasubject coefficient of variation (CV%) for the primary PK parameters has been reported as "
        "consistent with standard crossover BE study design."
    ))
    doc.add_paragraph()

    # ── SECTION 2: OBJECTIVES ─────────────────────────────────────────────────
    add_heading(doc, "2. STUDY OBJECTIVES", level=1)
    add_heading(doc, "2.1 Primary Objective", level=2)
    add_para(doc, (
        "To assess the bioequivalence of {{DRUG_NAME}} {{DOSE}} {{FORMULATION}} (Test) versus "
        "{{REFERENCE_PRODUCT}} (Reference) following a single oral dose under fasting conditions in "
        "healthy adult subjects."
    ))
    add_heading(doc, "2.2 Secondary Objectives", level=2)
    sec_obj = [
        "To assess the safety and tolerability of both formulations.",
        "To determine additional PK parameters: Tmax, Kel, t½.",
        "To confirm washout adequacy between study periods.",
    ]
    for obj in sec_obj:
        p = doc.add_paragraph(obj, style="List Bullet")
        p.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # ── SECTION 3: INVESTIGATIONAL PRODUCTS ───────────────────────────────────
    add_heading(doc, "3. INVESTIGATIONAL MEDICINAL PRODUCTS", level=1)
    add_heading(doc, "3.1 Test Product", level=2)

    test_table = doc.add_table(rows=6, cols=2)
    test_table.style = "Table Grid"
    test_fields = [
        ("Drug Name", "{{DRUG_NAME}}"),
        ("Dose", "{{DOSE}}"),
        ("Formulation", "{{FORMULATION}}"),
        ("Manufacturer / Sponsor", "{{SPONSOR_NAME}}, {{SPONSOR_COUNTRY}}"),
        ("Route of Administration", "Oral"),
        ("Storage Condition", "As per label — store below 30°C, protect from moisture"),
    ]
    for i, (label, value) in enumerate(test_fields):
        row = test_table.rows[i]
        lrun = row.cells[0].paragraphs[0].add_run(label)
        lrun.font.bold = True
        lrun.font.size = Pt(9)
        vrun = row.cells[1].paragraphs[0].add_run(value)
        vrun.font.size = Pt(9)
        if "{{" in value:
            vrun.font.color.rgb = TEAL
            vrun.font.italic = True

    doc.add_paragraph()
    add_heading(doc, "3.2 Reference Product", level=2)

    ref_table = doc.add_table(rows=4, cols=2)
    ref_table.style = "Table Grid"
    ref_fields = [
        ("Product Name", "{{REFERENCE_PRODUCT}}"),
        ("Country of Procurement", "{{REFERENCE_COUNTRY}}"),
        ("Dose / Strength", "{{DOSE}}"),
        ("Route of Administration", "Oral"),
    ]
    for i, (label, value) in enumerate(ref_fields):
        row = ref_table.rows[i]
        lrun = row.cells[0].paragraphs[0].add_run(label)
        lrun.font.bold = True
        lrun.font.size = Pt(9)
        vrun = row.cells[1].paragraphs[0].add_run(value)
        vrun.font.size = Pt(9)
        if "{{" in value:
            vrun.font.color.rgb = TEAL
            vrun.font.italic = True

    doc.add_paragraph()
    doc.add_page_break()

    # ── SECTION 4: STUDY DESIGN ───────────────────────────────────────────────
    add_heading(doc, "4. STUDY DESIGN", level=1)
    add_para(doc, (
        "This is a randomized, open-label, single-dose, two-period, two-sequence, crossover "
        "bioequivalence study in healthy adult subjects under fasting conditions. Subjects will be "
        "randomized to receive either the Test or Reference formulation in Period I, followed by the "
        "alternate formulation in Period II, with a washout period of {{WASHOUT_DAYS}} days between doses."
    ))

    add_heading(doc, "4.1 Washout Period", level=2)
    add_para(doc, (
        "A washout period of {{WASHOUT_DAYS}} days will be maintained between the two dosing periods. "
        "This corresponds to at least 5 elimination half-lives of {{DRUG_NAME}} (t½ ≈ {{HALF_LIFE}} hours), "
        "ensuring complete elimination of the drug before the subsequent dosing."
    ))

    add_heading(doc, "4.2 Confinement", level=2)
    add_para(doc, (
        "Subjects will be confined to the clinical unit for {{CONFINEMENT_HOURS}} hours post-dose per period. "
        "Ambulatory follow-up visits will be conducted at: {{AMBULATORY_VISITS}}."
    ))

    add_heading(doc, "4.3 Posture Restriction", level=2)
    add_para(doc, "{{POSTURE_RESTRICTION}}")

    add_heading(doc, "4.4 Dosing Conditions", level=2)
    add_para(doc, (
        "Subjects will fast for a minimum of 10 hours overnight prior to each dosing session. "
        "The study drug will be administered with 240 mL of water at ambient temperature. "
        "Water will be permitted ad libitum, except for 1 hour before and 2 hours after dosing. "
        "A standardized meal will be provided 4 hours post-dose."
    ))
    doc.add_paragraph()
    doc.add_page_break()

    # ── SECTION 5: SUBJECT SELECTION ──────────────────────────────────────────
    add_heading(doc, "5. SUBJECT SELECTION", level=1)
    add_heading(doc, "5.1 Inclusion Criteria", level=2)
    add_para(doc, "Subjects must meet ALL of the following criteria to be eligible:")
    inclusion = [
        "I1.  Healthy adult male or female subjects, aged 18 to 55 years (inclusive) at the time of screening.",
        "I2.  Body weight ≥50 kg (females) or ≥55 kg (males); Body Mass Index (BMI) between 18.5 and 30.0 kg/m² (inclusive).",
        "I3.  Non-smokers or ex-smokers who have ceased smoking for at least 6 months prior to screening.",
        "I4.  Willingness to use adequate contraception for females of childbearing potential; negative serum/urine pregnancy test at screening and check-in.",
        "I5.  Negative screening tests for HIV, HBsAg, and HCV antibody.",
        "I6.  Negative urine drugs-of-abuse screen and alcohol breath test at screening and each check-in.",
        "I7.  Able to understand and sign the written Informed Consent Form (ICF) prior to any study-related procedures.",
        "I8.  Willing and able to comply with all study procedures and visit schedules.",
        "I9.  Normal 12-lead ECG with QTcF ≤450 ms (males) or ≤470 ms (females) at screening.",
        "I10. All clinical laboratory values within the normal range or assessed as not clinically significant by the Investigator.",
    ]
    for ic in inclusion:
        p = doc.add_paragraph(ic, style="List Bullet")
        p.runs[0].font.size = Pt(9)

    doc.add_paragraph()
    add_heading(doc, "5.2 Exclusion Criteria", level=2)
    add_para(doc, "Subjects will be excluded if ANY of the following criteria are met:")
    exclusion = [
        "E1.  Any significant medical history, current medical condition, or finding on physical examination that, in the Investigator's opinion, could interfere with the study or put the subject at risk.",
        "E2.  History of hypersensitivity or allergy to {{DRUG_NAME}}, related drug classes, or any excipients of the study formulations.",
        "E3.  Use of any prescription or non-prescription medication, vitamins, or herbal supplements within 14 days (or 5 half-lives, whichever is longer) prior to first dosing.",
        "E4.  Participation in another clinical study within 90 days prior to first study drug administration.",
        "E5.  Donation of blood (>450 mL) or blood products within 90 days prior to screening.",
        "E6.  History of drug or alcohol abuse within the past 2 years.",
        "E7.  Positive pregnancy test or breastfeeding (females).",
        "E8.  Presence of any gastrointestinal, hepatic, or renal disease that may affect drug absorption or elimination.",
        "E9.  Consumption of grapefruit, grapefruit juice, or Seville orange within 14 days prior to first dosing.",
        "E10. Use of any enzyme-inducing or enzyme-inhibiting drug within 30 days prior to first dosing.",
    ]
    for ec in exclusion:
        p = doc.add_paragraph(ec, style="List Bullet")
        p.runs[0].font.size = Pt(9)

    doc.add_paragraph()
    doc.add_page_break()

    # ── SECTION 6: STUDY PROCEDURES ───────────────────────────────────────────
    add_heading(doc, "6. STUDY PROCEDURES", level=1)
    add_heading(doc, "6.1 Screening (Day -28 to Day -2)", level=2)
    screening_items = [
        "Informed consent obtained",
        "Medical history and physical examination",
        "12-lead ECG",
        "Vital signs (BP, HR, temperature, respiratory rate)",
        "Clinical laboratory tests (haematology, biochemistry, urinalysis)",
        "Serology (HIV, HBsAg, HCV Ab)",
        "Drugs-of-abuse screen and alcohol breath test",
        "Pregnancy test (females of childbearing potential)",
    ]
    for item in screening_items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(9)

    add_heading(doc, "6.2 Check-in (Day -1 of each period)", level=2)
    checkin = [
        "Physical examination and vital signs",
        "Drugs-of-abuse screen and alcohol breath test",
        "Pregnancy test (females)",
        "Clinical laboratory safety tests",
        "12-lead ECG",
        "Confirm fasting compliance (≥10 hours prior to dosing)",
    ]
    for item in checkin:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(9)

    add_heading(doc, "6.3 Dosing and Post-Dose (Day 1 of each period)", level=2)
    add_para(doc, (
        "A single oral dose of {{DOSE}} {{DRUG_NAME}} (Test) or {{REFERENCE_PRODUCT}} (Reference) "
        "will be administered with 240 mL water. {{POSTURE_RESTRICTION}}. "
        "Serial blood samples will be collected according to the PK sampling schedule in Section 7."
    ))

    add_heading(doc, "6.4 Safety Follow-up", level=2)
    add_para(doc, (
        "A follow-up telephone call or clinic visit will be conducted within 7 days of the last "
        "blood sample collection to assess any delayed adverse events."
    ))
    doc.add_paragraph()
    doc.add_page_break()

    # ── SECTION 7: PK SAMPLING ────────────────────────────────────────────────
    add_heading(doc, "7. PHARMACOKINETIC BLOOD SAMPLING", level=1)
    add_heading(doc, "7.1 Sampling Schedule", level=2)
    add_para(doc, (
        "Serial blood samples will be collected at the following time points post-dose (hours):"
    ))
    p = doc.add_paragraph()
    run = p.add_run("{{PK_SAMPLING_SCHEDULE}}")
    run.font.color.rgb = TEAL
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.bold = True

    doc.add_paragraph()
    add_heading(doc, "7.2 Blood Collection Details", level=2)
    blood_table = doc.add_table(rows=6, cols=2)
    blood_table.style = "Table Grid"
    blood_fields = [
        ("Collection tube", "K₂EDTA vacutainer (lavender cap)"),
        ("Volume per sample", "3 mL whole blood"),
        ("Total volume (per period)", "Calculated based on number of timepoints"),
        ("Centrifugation", "3000 rpm, 10 min, 4°C within 60 minutes of collection"),
        ("Aliquots", "Two aliquots per sample; store at −70°C until analysis"),
        ("Sample labelling", "Subject ID, period, nominal time, actual time, date"),
    ]
    for i, (label, value) in enumerate(blood_fields):
        row = blood_table.rows[i]
        lrun = row.cells[0].paragraphs[0].add_run(label)
        lrun.font.bold = True
        lrun.font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run(value).font.size = Pt(9)

    doc.add_paragraph()
    doc.add_page_break()

    # ── SECTION 8: PK ANALYSIS ────────────────────────────────────────────────
    add_heading(doc, "8. PHARMACOKINETIC ANALYSIS", level=1)
    add_para(doc, (
        "Plasma concentrations of {{DRUG_NAME}} will be measured by a validated bioanalytical method "
        "(LC-MS/MS). Non-compartmental analysis (NCA) will be performed using Phoenix WinNonlin® or equivalent."
    ))
    pk_params = [
        ("AUC₀‑t", "Area under the plasma concentration-time curve from time 0 to last measurable concentration"),
        ("AUC₀‑∞", "AUC₀‑t + Clast/Kel"),
        ("Cmax", "Maximum observed plasma concentration"),
        ("Tmax", "Time to Cmax (observed)"),
        ("t½", "Terminal elimination half-life = ln(2)/Kel"),
        ("Kel", "Terminal elimination rate constant (log-linear regression)"),
    ]
    pk_table = doc.add_table(rows=len(pk_params), cols=2)
    pk_table.style = "Table Grid"
    for i, (param, desc) in enumerate(pk_params):
        row = pk_table.rows[i]
        row.cells[0].paragraphs[0].add_run(param).font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run(desc).font.size = Pt(9)
    doc.add_paragraph()
    doc.add_page_break()

    # ── SECTION 9: STATISTICAL ANALYSIS ───────────────────────────────────────
    add_heading(doc, "9. STATISTICAL ANALYSIS", level=1)
    add_heading(doc, "9.1 Sample Size Justification", level=2)
    p = doc.add_paragraph()
    run = p.add_run("{{SAMPLE_SIZE_BASIS}}")
    run.font.color.rgb = TEAL
    run.font.italic = True
    run.font.size = Pt(10)

    add_heading(doc, "9.2 Bioequivalence Analysis", level=2)
    add_para(doc, (
        "Log-transformed AUC₀‑t, AUC₀‑∞, and Cmax will be analyzed using a mixed-effects ANOVA model "
        "with sequence, period, and treatment as fixed effects, and subject nested within sequence as a "
        "random effect. The 90% confidence intervals (CI) for the test/reference geometric mean ratio (GMR) "
        "will be constructed. Bioequivalence is concluded if both the lower and upper bounds of the 90% CI "
        "fall within the acceptance range of 80.00%–125.00%."
    ))
    doc.add_paragraph()
    doc.add_page_break()

    # ── SECTION 10: SAFETY ────────────────────────────────────────────────────
    add_heading(doc, "10. SAFETY MONITORING", level=1)
    add_para(doc, (
        "Safety will be assessed throughout the study by monitoring adverse events (AEs), vital signs, "
        "12-lead ECG, physical examination, and clinical laboratory tests. All AEs will be coded using "
        "MedDRA and graded per CTCAE v5.0."
    ))
    add_heading(doc, "10.1 Drug-Specific Safety Requirements", level=2)
    p = doc.add_paragraph()
    run = p.add_run("[Drug-specific safety monitoring flags will be inserted here based on the drug profile]")
    run.font.color.rgb = TEAL
    run.font.italic = True
    run.font.size = Pt(10)
    doc.add_paragraph()

    # ── SECTION 11: DATA MANAGEMENT ───────────────────────────────────────────
    add_heading(doc, "11. DATA MANAGEMENT", level=1)
    add_para(doc, (
        f"All data will be collected on validated electronic case report forms (eCRFs) in compliance with "
        f"ICH E6(R3) GCP guidelines. Source documents will be retained at {COMPANY} for a minimum of 15 years. "
        f"Data will be locked following quality review and reconciliation prior to statistical analysis."
    ))
    doc.add_paragraph()

    # ── SECTION 12: ETHICS ────────────────────────────────────────────────────
    add_heading(doc, "12. ETHICS AND REGULATORY COMPLIANCE", level=1)
    add_para(doc, (
        f"This study will be conducted in accordance with the Declaration of Helsinki (2013), ICH E6(R3) "
        f"GCP, and applicable local regulations. The protocol will be submitted for approval to an "
        f"Independent Ethics Committee (IEC) prior to initiation. All subjects will provide written informed "
        f"consent before any study procedures. This study will be registered on the Clinical Trials Registry "
        f"of India (CTRI) prior to enrolment. Regulatory submission targets: {{REGULATORY_TARGETS}}."
    ))
    doc.add_paragraph()

    # ── SECTION 13: QA ────────────────────────────────────────────────────────
    add_heading(doc, "13. QUALITY ASSURANCE", level=1)
    add_para(doc, (
        f"The study will be monitored by {COMPANY} Quality Assurance personnel. Sponsor monitors will "
        f"conduct periodic site visits. A final audit will be performed prior to database lock. "
        f"All deviations from this protocol will be documented and reported."
    ))
    doc.add_paragraph()

    # ── SIGNATURES ────────────────────────────────────────────────────────────
    add_heading(doc, "PROTOCOL SIGNATURES", level=1)
    sig_table = doc.add_table(rows=4, cols=3)
    sig_table.style = "Table Grid"
    sig_headers = ["Role", "Name & Signature", "Date"]
    header_row = sig_table.rows[0]
    for i, h in enumerate(sig_headers):
        set_cell_bg(header_row.cells[i], "1E2A3A")
        run = header_row.cells[i].paragraphs[0].add_run(h)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(9)
    sig_roles = ["Principal Investigator", "Sponsor Medical Monitor", "Biostatistician"]
    for i, role in enumerate(sig_roles):
        row = sig_table.rows[i + 1]
        row.cells[0].paragraphs[0].add_run(role).font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run(" ").font.size = Pt(9)
        row.cells[2].paragraphs[0].add_run(" ").font.size = Pt(9)

    doc.add_page_break()

    # ── APPENDIX A: SCHEDULE OF EVENTS ────────────────────────────────────────
    add_heading(doc, "APPENDIX A: SCHEDULE OF EVENTS", level=1)
    add_para(doc, "The table below summarizes all assessments by visit and period.", italic=True)

    soe_cols = ["Assessment", "Screening\n(D-28 to D-2)", "Check-in\n(D-1)", "Dosing Day\n(D1)", "Discharge\n(D2)", "Ambulatory\nVisits", "Follow-up\n(D7±2)"]
    soe_table = doc.add_table(rows=13, cols=len(soe_cols))
    soe_table.style = "Table Grid"
    header_row = soe_table.rows[0]
    for i, col in enumerate(soe_cols):
        set_cell_bg(header_row.cells[i], "1E2A3A")
        run = header_row.cells[i].paragraphs[0].add_run(col)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(8)

    soe_data = [
        ("Informed Consent", "X", "", "", "", "", ""),
        ("Medical History", "X", "", "", "", "", ""),
        ("Physical Exam", "X", "X", "", "X", "", "X"),
        ("Vital Signs", "X", "X", "X*", "X", "", "X"),
        ("12-lead ECG", "X", "X", "X†", "", "", "X"),
        ("Lab Safety Tests", "X", "X", "", "X", "", "X"),
        ("Drugs of Abuse Screen", "X", "X", "", "", "", ""),
        ("Pregnancy Test (F)", "X", "X", "", "", "", ""),
        ("Study Drug Admin.", "", "", "X", "", "", ""),
        ("PK Blood Sampling", "", "", "{{PK_SAMPLING_SCHEDULE}}", "", "", ""),
        ("Adverse Event Assessment", "", "X", "X", "X", "X", "X"),
        ("Concomitant Medication", "X", "X", "X", "X", "X", "X"),
    ]
    for i, (label, *cols) in enumerate(soe_data):
        row = soe_table.rows[i + 1]
        row.cells[0].paragraphs[0].add_run(label).font.size = Pt(8)
        for j, val in enumerate(cols):
            run = row.cells[j + 1].paragraphs[0].add_run(val)
            run.font.size = Pt(8)
            if "{{" in val:
                run.font.color.rgb = TEAL
                run.font.italic = True

    doc.add_paragraph()
    add_para(doc, "* Vital signs: pre-dose and at 1h, 2h, 4h, 8h, 12h, 24h post-dose", size=8, italic=True, color=RGBColor(0x71, 0x80, 0x96))
    add_para(doc, "† ECG: pre-dose and at 3h, 8h post-dose (or per drug-specific requirements)", size=8, italic=True, color=RGBColor(0x71, 0x80, 0x96))
    add_para(doc, "{{AMBULATORY_VISITS}} — ambulatory blood sampling visits (subjects return to clinic for PK samples)", size=8, italic=True, color=TEAL)

    doc.add_page_break()

    # ── FOOTER NOTE ───────────────────────────────────────────────────────────
    add_heading(doc, "DOCUMENT CONTROL", level=1)
    dc_table = doc.add_table(rows=4, cols=2)
    dc_table.style = "Table Grid"
    dc_data = [
        ("Document Type", "Study Protocol"),
        ("Organisation", COMPANY),
        ("Template Version", "v1.0"),
        ("Generated by", "TrialOS Protocol AI — {{DATE}}"),
    ]
    for i, (label, value) in enumerate(dc_data):
        row = dc_table.rows[i]
        lrun = row.cells[0].paragraphs[0].add_run(label)
        lrun.font.bold = True
        lrun.font.size = Pt(9)
        vrun = row.cells[1].paragraphs[0].add_run(value)
        vrun.font.size = Pt(9)
        if "{{" in value:
            vrun.font.color.rgb = TEAL
            vrun.font.italic = True

    doc.save(OUTPUT_PATH)
    print(f"✓ Template saved: {OUTPUT_PATH}")
    size_kb = os.path.getsize(OUTPUT_PATH) // 1024
    print(f"  File size: {size_kb} KB")
    print(f"  Company: {COMPANY}")
    print(f"  Placeholders: {{DRUG_NAME}}, {{DOSE}}, {{FORMULATION}}, {{REFERENCE_PRODUCT}},")
    print(f"                {{REFERENCE_COUNTRY}}, {{SPONSOR_NAME}}, {{SPONSOR_COUNTRY}},")
    print(f"                {{REGULATORY_TARGETS}}, {{TARGET_SUBJECTS}}, {{SAMPLE_SIZE_RECOMMENDED}},")
    print(f"                {{HALF_LIFE}}, {{TMAX}}, {{WASHOUT_DAYS}}, {{CONFINEMENT_HOURS}},")
    print(f"                {{PK_SAMPLING_SCHEDULE}}, {{AMBULATORY_VISITS}}, {{POSTURE_RESTRICTION}},")
    print(f"                {{SAMPLE_SIZE_BASIS}}, {{STUDY_ID}}, {{DATE}}, {{VERSION}}")


if __name__ == "__main__":
    build_template()
